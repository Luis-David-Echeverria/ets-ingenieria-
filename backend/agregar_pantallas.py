"""Agrega la seccion de pantallas al documento existente (sin perder cambios).

Uso: python agregar_pantallas.py
Abre ../docs/Documentacion.docx, agrega la seccion de pantallas con las
capturas y las descripciones, y guarda el mismo archivo.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

DOCS = Path(__file__).resolve().parent.parent / "docs"
PANT = DOCS / "imagenes" / "pantallas"
DOC_PATH = DOCS / "Documentacion.docx"

# Guinda institucional del IPN
GUINDA = RGBColor(0x7B, 0x2C, 0x3B)
GRIS = RGBColor(0x44, 0x44, 0x44)

# Cada pantalla: archivo, titulo y descripcion documentada
PANTALLAS = [
    ("portada.png", "Portada",
     "Pantalla inicial del sistema. Muestra la portada del cuento con la "
     "ilustración de los animales. El niño pasa la página para comenzar."),
    ("identificar.png", "Identificar animal",
     "Sección de identificación. El niño toca el recuadro para subir una foto "
     "de un animal desde su dispositivo. La imagen se analiza automáticamente."),
    ("identificado.png", "Animal identificado",
     "Resultado de la identificación. El sistema muestra el ícono y el nombre "
     "de la especie predicha junto con el porcentaje de confianza, y lo "
     "reproduce en voz alta."),
    ("no-identificado.png", "Animal no reconocido",
     "Cuando la confianza del modelo es menor al 85%, el sistema informa que "
     "no pudo reconocer al animal y sugiere probar con otra foto."),
    ("quiz-neutral.png", "Quiz — pregunta",
     "Modo de juego. El sistema muestra el ícono de un animal y tres nombres "
     "como opciones. El niño debe adivinar de qué especie se trata."),
    ("quiz-respondido.png", "Quiz — respuesta",
     "Retroalimentación del quiz. Al elegir una opción, el sistema la marca en "
     "verde si es correcta o en rojo si es incorrecta, revela el nombre correcto "
     "y lo dice en voz alta."),
    ("contraportada.png", "Contraportada",
     "Pantalla final. Muestra el puntaje obtenido en el quiz mediante estrellas "
     "y los créditos del autor. Al volver a la portada, el quiz se reinicia."),
]


def main():
    doc = Document(DOC_PATH)

    # Nueva pagina y titulo de la seccion
    doc.add_page_break()
    h = doc.add_heading(level=1)
    run = h.add_run("12. Pantallas del Sistema")
    run.font.color.rgb = GUINDA

    p = doc.add_paragraph()
    p.add_run(
        "A continuación se documentan las pantallas principales del sistema, "
        "que corresponden a las trayectorias de los casos de uso."
    )

    figura = 9  # continua la numeracion de figuras del documento
    for archivo, nombre, descripcion in PANTALLAS:
        ruta = PANT / archivo
        if not ruta.exists():
            continue
        # Subtitulo de la pantalla
        sub = doc.add_heading(level=2)
        r = sub.add_run(nombre)
        r.font.color.rgb = GUINDA
        # Descripcion
        doc.add_paragraph(descripcion)
        # Imagen centrada
        doc.add_picture(str(ruta), width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Pie de figura
        pie = doc.add_paragraph()
        pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp = pie.add_run(f"Figura {figura}. Pantalla: {nombre}")
        rp.italic = True
        rp.font.size = Pt(9)
        rp.font.color.rgb = GRIS
        doc.add_paragraph()
        figura += 1

    doc.save(DOC_PATH)
    print("Pantallas agregadas a", DOC_PATH)


if __name__ == "__main__":
    main()
