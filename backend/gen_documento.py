"""Genera el documento tecnico del proyecto en formato .docx.

Uso: python gen_documento.py
Produce ../docs/Documentacion.docx con portada, indice, secciones y los
diagramas insertados. Usa el color guinda institucional del IPN.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

DOCS = Path(__file__).resolve().parent.parent / "docs"
IMG = DOCS / "imagenes"

# Guinda institucional del IPN
GUINDA = RGBColor(0x7B, 0x2C, 0x3B)
GRIS = RGBColor(0x44, 0x44, 0x44)


def titulo(doc, texto, nivel=1):
    """Agregar un encabezado con el color guinda."""
    h = doc.add_heading(level=nivel)
    run = h.add_run(texto)
    run.font.color.rgb = GUINDA
    return h


def parrafo(doc, texto, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.italic = italic
    return p


def tabla(doc, encabezados, filas):
    """Crear una tabla con estilo y encabezado guinda."""
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, texto in enumerate(encabezados):
        hdr[i].paragraphs[0].add_run(texto).bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, valor in enumerate(fila):
            celdas[i].text = str(valor)
    doc.add_paragraph()
    return t


def imagen(doc, nombre, ancho=6.0, pie=None):
    """Insertar una imagen centrada con pie opcional."""
    ruta = IMG / nombre
    if ruta.exists():
        doc.add_picture(str(ruta), width=Inches(ancho))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if pie:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(pie)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRIS
    doc.add_paragraph()


def portada(doc):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INSTITUTO POLITÉCNICO NACIONAL")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = GUINDA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Escuela Superior de Cómputo")
    run.font.size = Pt(13)
    run.font.color.rgb = GUINDA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ingeniería en Inteligencia Artificial")
    run.font.size = Pt(12)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("El Gran Libro de los Animales")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = GUINDA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de clasificación de especies de animales para niños")
    run.italic = True
    run.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()

    datos = [
        ("Alumno", "Luis David Echeverría Pérez"),
        ("Boleta", "2022630724"),
        ("Materia", "Ingeniería de Software para Sistemas Inteligentes"),
        ("Profesora", "Idalia Maldonado Castillo"),
        ("Evaluación", "Examen a Título de Suficiencia (ETS)"),
    ]
    for etiqueta, valor in datos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{etiqueta}: ")
        run.bold = True
        p.add_run(valor)

    doc.add_page_break()


def seccion_indice(doc):
    titulo(doc, "Índice", 1)
    secciones = [
        "1. Alcance del proyecto",
        "2. Objetivos",
        "3. Análisis de Requerimientos",
        "4. Reglas de Negocio",
        "5. Diagrama de Clases",
        "6. Diagramas de Casos de Uso",
        "7. Diagramas de Secuencia",
        "8. Máquina de Estados",
        "9. Arquitectura del Sistema",
        "10. Evaluación del Modelo (Métricas)",
        "11. Pruebas de Aceptación",
    ]
    for s in secciones:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_page_break()


def seccion_alcance(doc):
    titulo(doc, "1. Alcance del proyecto", 1)
    parrafo(doc,
        "El presente proyecto consiste en un sistema web educativo capaz de "
        "identificar automáticamente especies de animales a partir de imágenes, "
        "utilizando técnicas de procesamiento de imágenes y aprendizaje profundo. "
        "El sistema está dirigido a niños de primaria y a sus profesores, con el "
        "propósito de que los niños aprendan a distinguir entre distintas especies "
        "de animales de una manera lúdica e interactiva.")

    titulo(doc, "Incluye", 2)
    incluye = [
        "Clasificación de al menos 10 especies de animales mediante un modelo de "
        "aprendizaje profundo (transfer learning con MobileNetV2).",
        "Interfaz gráfica amigable con estética de cuento infantil.",
        "Módulo de identificación: el niño sube una foto y recibe la predicción "
        "de la especie, que se reproduce en voz alta.",
        "Módulo de quiz: el sistema muestra un animal y el niño adivina su nombre "
        "en opción múltiple, con retroalimentación inmediata.",
        "Preprocesamiento automático de imágenes (redimensionamiento y normalización).",
        "Validaciones de entrada tanto en el cliente como en el servidor.",
        "Evaluación del modelo con métricas de clasificación (precisión, "
        "sensibilidad y especificidad).",
    ]
    for i in incluye:
        doc.add_paragraph(i, style="List Bullet")

    titulo(doc, "No incluye", 2)
    no_incluye = [
        "Registro de usuarios ni autenticación (login).",
        "Persistencia del progreso del niño a largo plazo (base de datos).",
        "Aplicación móvil nativa (el sistema es web).",
        "Reconocimiento de especies fuera de las 10 entrenadas.",
        "Detección de múltiples animales en una misma imagen.",
    ]
    for i in no_incluye:
        doc.add_paragraph(i, style="List Bullet")
    doc.add_page_break()


def seccion_objetivos(doc):
    titulo(doc, "2. Objetivos", 1)
    titulo(doc, "Objetivo general", 2)
    parrafo(doc,
        "Desarrollar un sistema inteligente que clasifique automáticamente al menos "
        "10 especies de animales a partir de imágenes, con una interfaz amigable "
        "para niños de primaria, alcanzando una buena eficiencia y capacidad de "
        "generalización del modelo.")
    titulo(doc, "Objetivos particulares", 2)
    objs = [
        "Entrenar un clasificador mediante transfer learning con una precisión "
        "superior al 85% en el conjunto de prueba.",
        "Preprocesar las imágenes (redimensionar a 224x224 y normalizar) para "
        "homogeneizar la entrada al modelo.",
        "Exponer el modelo a través de una API REST (FastAPI).",
        "Construir una interfaz interactiva con un modo de identificación y un "
        "modo de quiz.",
        "Reproducir en voz alta el nombre de la especie para reforzar el aprendizaje.",
        "Evaluar el modelo con métricas de precisión, sensibilidad y especificidad.",
    ]
    for o in objs:
        doc.add_paragraph(o, style="List Bullet")
    doc.add_page_break()


def seccion_requerimientos(doc):
    titulo(doc, "3. Análisis de Requerimientos", 1)
    titulo(doc, "Requerimientos funcionales", 2)
    tabla(doc, ["ID", "Requerimiento"], [
        ["RF-01", "Cargar una imagen y clasificar la especie del animal"],
        ["RF-02", "Devolver la especie predicha junto con su nivel de confianza"],
        ["RF-03", "Reproducir el nombre de la especie en voz alta"],
        ["RF-04", "Mostrar un animal aleatorio en el modo quiz"],
        ["RF-05", "Presentar opción múltiple con tres nombres de especie"],
        ["RF-06", "Validar la respuesta del quiz e indicar si es correcta"],
        ["RF-07", "Preprocesar la imagen (redimensionar y normalizar)"],
        ["RF-08", "Mostrar el puntaje final del quiz en estrellas"],
    ])
    titulo(doc, "Requerimientos no funcionales", 2)
    tabla(doc, ["ID", "Requerimiento"], [
        ["RNF-01", "Interfaz amigable e intuitiva para niños de primaria"],
        ["RNF-02", "Tiempo de predicción menor a 3 segundos"],
        ["RNF-03", "Precisión (accuracy) del modelo mayor o igual a 85%"],
        ["RNF-04", "Ejecutable en local (frontend y backend separados)"],
        ["RNF-05", "El modelo debe generalizar y evitar el sobreajuste"],
        ["RNF-06", "Validación de entradas en cliente y servidor"],
    ])
    doc.add_page_break()


def seccion_reglas(doc):
    titulo(doc, "4. Reglas de Negocio", 1)
    tabla(doc, ["ID", "Regla"], [
        ["RN-01", "El sistema solo reconoce las 10 especies entrenadas; si la "
                  "confianza es menor a 85%, informa que no pudo reconocer al animal."],
        ["RN-02", "El quiz siempre incluye la respuesta correcta entre las tres opciones."],
        ["RN-03", "Los distractores se eligen aleatoriamente entre las demás especies."],
        ["RN-04", "En una ronda de quiz ningún animal se repite."],
        ["RN-05", "Solo se aceptan archivos de imagen válidos y de hasta 10 MB."],
    ])
    doc.add_page_break()


def seccion_diagramas(doc):
    titulo(doc, "5. Diagrama de Clases", 1)
    parrafo(doc,
        "Representa las clases y servicios principales del backend y los servicios "
        "del modelo de inteligencia artificial.")
    imagen(doc, "clases.png", 6.0, "Figura 1. Diagrama de clases")
    doc.add_page_break()

    titulo(doc, "6. Diagramas de Casos de Uso", 1)
    parrafo(doc,
        "Los actores del sistema son el Niño (usuario principal) y el Profesor "
        "(guía y supervisor).")
    imagen(doc, "casos-de-uso.png", 6.5, "Figura 2. Diagrama de casos de uso")
    doc.add_page_break()

    titulo(doc, "7. Diagramas de Secuencia", 1)
    parrafo(doc, "Flujo de la identificación de un animal por foto:")
    imagen(doc, "secuencia-identificar.png", 4.5, "Figura 3. Secuencia: identificar animal")
    parrafo(doc, "Flujo de una pregunta del quiz:")
    imagen(doc, "secuencia-quiz.png", 4.5, "Figura 4. Secuencia: responder quiz")
    doc.add_page_break()

    titulo(doc, "8. Máquina de Estados", 1)
    parrafo(doc, "Estados de navegación del libro (recorrido del cuento):")
    imagen(doc, "maquina-estados-libro.png", 6.5, "Figura 5. Máquina de estados del libro")
    parrafo(doc, "Estados de una pregunta del quiz:")
    imagen(doc, "maquina-estados-quiz.png", 6.0, "Figura 6. Máquina de estados del quiz")
    doc.add_page_break()

    titulo(doc, "9. Arquitectura del Sistema", 1)
    parrafo(doc,
        "El sistema sigue una arquitectura cliente-servidor de tres capas: un "
        "frontend en React que consume una API REST en FastAPI, la cual utiliza "
        "el modelo MobileNetV2 entrenado mediante transfer learning.")
    imagen(doc, "arquitectura.png", 5.5, "Figura 7. Arquitectura general")
    doc.add_page_break()


def seccion_metricas(doc):
    titulo(doc, "10. Evaluación del Modelo (Métricas)", 1)
    parrafo(doc,
        "El modelo se evaluó sobre un conjunto de validación de 5,235 imágenes "
        "(el 20% del dataset, no vistas durante el entrenamiento). Se obtuvo una "
        "exactitud global (accuracy) de 94.63%.")
    tabla(doc, ["Especie", "Precisión", "Sensibilidad", "Especificidad", "F1"], [
        ["Perro", "0.920", "0.968", "0.981", "0.943"],
        ["Caballo", "0.957", "0.928", "0.995", "0.942"],
        ["Elefante", "0.953", "0.964", "0.997", "0.958"],
        ["Mariposa", "0.968", "0.942", "0.998", "0.955"],
        ["Gallina", "0.943", "0.965", "0.992", "0.954"],
        ["Gato", "0.993", "0.859", "1.000", "0.921"],
        ["Vaca", "0.935", "0.848", "0.995", "0.889"],
        ["Oveja", "0.863", "0.944", "0.988", "0.902"],
        ["Araña", "0.977", "0.987", "0.994", "0.982"],
        ["Ardilla", "0.973", "0.939", "0.998", "0.956"],
        ["Promedio", "0.948", "0.934", "0.994", "0.940"],
    ])
    parrafo(doc, "La matriz de confusión muestra el detalle de aciertos y errores por clase:")
    imagen(doc, "matriz-confusion.png", 6.0, "Figura 8. Matriz de confusión (validación)")
    doc.add_page_break()


def seccion_pruebas(doc):
    titulo(doc, "11. Pruebas de Aceptación", 1)
    titulo(doc, "Pruebas funcionales", 2)
    tabla(doc, ["ID", "Caso de prueba", "Resultado esperado", "Estado"], [
        ["PA-01", "Identificar un perro", "Predice Perro con alta confianza", "Pasó"],
        ["PA-02", "Identificar un elefante", "Predice Elefante", "Pasó"],
        ["PA-03", "Identificar una araña", "Predice Araña", "Pasó"],
        ["PA-04", "Umbral de confianza", "Muestra 'no reconocido' si < 85%", "Pasó"],
        ["PA-05", "Reproducción de voz", "Dice el nombre en voz alta", "Pasó"],
        ["PA-06", "Generar pregunta de quiz", "Muestra animal y 3 opciones", "Pasó"],
        ["PA-07", "Responder correctamente", "Marca verde y suma estrella", "Pasó"],
        ["PA-08", "Responder incorrectamente", "Marca rojo y revela el correcto", "Pasó"],
        ["PA-09", "Animales sin repetir", "Ningún animal se repite en la ronda", "Pasó"],
        ["PA-10", "Puntaje final", "Muestra estrellas según aciertos", "Pasó"],
    ])
    titulo(doc, "Pruebas de validación (robustez)", 2)
    tabla(doc, ["ID", "Entrada", "Resultado esperado", "Estado"], [
        ["PV-01", "Archivo no-imagen (cliente)", "Rechaza y avisa al usuario", "Pasó"],
        ["PV-02", "Archivo no-imagen (servidor)", "HTTP 400", "Pasó"],
        ["PV-03", "Imagen mayor a 10 MB", "Rechaza (HTTP 413)", "Pasó"],
        ["PV-04", "Imagen corrupta", "HTTP 400 imagen no válida", "Pasó"],
        ["PV-05", "Especie desconocida en quiz", "HTTP 400", "Pasó"],
        ["PV-06", "Fallo del servidor", "Mensaje de error sin colgarse", "Pasó"],
    ])
    titulo(doc, "Rendimiento del modelo", 2)
    tabla(doc, ["Métrica", "Valor", "Requisito"], [
        ["Accuracy (validación)", "94.63%", "Mayor o igual a 85%"],
        ["Precisión (promedio)", "94.8%", "Requerida"],
        ["Sensibilidad (recall)", "93.4%", "Requerida"],
        ["Especificidad", "99.4%", "Requerida"],
        ["Tiempo de predicción", "Menor a 2 s", "Fluido"],
    ])


def main():
    doc = Document()
    # Fuente base
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    portada(doc)
    seccion_indice(doc)
    seccion_alcance(doc)
    seccion_objetivos(doc)
    seccion_requerimientos(doc)
    seccion_reglas(doc)
    seccion_diagramas(doc)
    seccion_metricas(doc)
    seccion_pruebas(doc)

    salida = DOCS / "Documentacion.docx"
    doc.save(salida)
    print("Documento generado en", salida)


if __name__ == "__main__":
    main()
